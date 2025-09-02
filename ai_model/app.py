import io
import re
import pdfplumber
import docx
import pytesseract
import fitz
from PIL import Image
from transformers import pipeline
from flask import Flask, request, jsonify, send_file
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet

# Flask app
app = Flask(__name__)

_classifier = None
def get_classifier():
    global _classifier
    if _classifier is None:
        _classifier = pipeline(
            "zero-shot-classification",
            model="valhalla/distilbart-mnli-12-3",
            device=-1  # CPU
        )
    return _classifier

# Section cues to check for agreements

POSITIVE_LABELS = [
    "agreement", "legal contract", "rental agreement", "lease agreement",
    "service agreement", "tenant-landlord agreement", "terms and conditions",
    "offer letter", "internship agreement", "employment contract",
    "student agreement", "job offer", "internship terms"
]


SECTION_CUES = [
    "agreement", "security deposit", "rental period", "payment terms",
    "termination", "arbitration", "jurisdiction",
    "witness", "signatory", "governing law", "parties", "definitions",
    "probation period", "internship duration", "performance", 
    "salary", "compensation", "notice period", "work expectations",
    "attendance", "leaves", "certificate", "offer letter"
]

# Helpers
def safe_join_text(parts):
    return "\n".join([p for p in parts if p])

def chunk_text(text, max_words=300, max_chunks=10):
    words = text.split()
    chunks = [" ".join(words[i:i + max_words]) for i in range(0, len(words), max_words)]
    return chunks[:max_chunks]

def heuristic_score(text):
    t = (text or "").lower()
    found = sum(1 for k in SECTION_CUES if re.search(r"\b" + re.escape(k) + r"\b", t))
    return found / max(1, len(SECTION_CUES))

def classify_agreement(text):
    details = {
        "chunks": 0,
        "votes": 0,
        "vote_ratio": 0.0,
        "heuristic": 0.0,
        "avg_chunk_score": 0.0,
        "reason": ""
    }
    if not text.strip():
        details["reason"] = "empty_text"
        return False, details
    chunks = chunk_text(text, max_words=300, max_chunks=10)
    details["chunks"] = len(chunks)
    votes, per_chunk_scores = 0, []
    CHUNK_THRESHOLD = 0.5
    for ch in chunks:
        try:
            classifier = get_classifier()
            res = classifier(
                ch,
                candidate_labels=POSITIVE_LABELS,
                multi_label=True,
                hypothesis_template="This text is about {}."
            )
            best = max(res["scores"]) if res["scores"] else 0.0
        except Exception as e:
            print(f"Model error: {e}")
            best = 0.0
        per_chunk_scores.append(best)
        if best >= CHUNK_THRESHOLD:
            votes += 1
    ratio = votes / len(chunks)
    heur = heuristic_score(text)
    details.update({
        "votes": votes,
        "vote_ratio": round(ratio, 3),
        "heuristic": round(heur, 3),
        "avg_chunk_score": round(sum(per_chunk_scores) / max(1, len(per_chunk_scores)), 3)
    })
    accept = (ratio >= 0.4) or (heur >= 0.4)
    if not accept:
        details["reason"] = "low_confidence"
    return accept, details

# File extraction
def extract_pdf(file_stream):
    try:
        file_stream.seek(0)
        with pdfplumber.open(file_stream) as pdf:
            return safe_join_text([p.extract_text() for p in pdf.pages])
    except Exception as e:
        print(f"PDF extract error: {e}")
        try:
            file_stream.seek(0)
            doc = fitz.open(stream=file_stream.read(), filetype="pdf")
            texts = []
            for p in doc:
                pix = p.get_pixmap(dpi=200)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                texts.append(pytesseract.image_to_string(img))
            return "\n".join(texts)
        except Exception as e2:
            print(f"PDF OCR error: {e2}")
            return ""

def extract_docx(file_stream):
    try:
        file_stream.seek(0)
        doc = docx.Document(io.BytesIO(file_stream.read()))
        return "\n".join(p.text for p in doc.paragraphs if p.text)
    except Exception as e:
        print(f"DOCX extract error: {e}")
        return ""

def extract_image(file_stream):
    try:
        file_stream.seek(0)
        img = Image.open(file_stream).convert("RGB")
        return pytesseract.image_to_string(img)
    except Exception as e:
        print(f"Image extract error: {e}")
        return ""

# Routes
@app.route("/active", methods=["GET"])
def active():
    return "active"

@app.route("/uploads", methods=["POST"])
def api_upload():
    if "file" not in request.files:
        return jsonify({"error": "No file uploaded"}), 400
    file = request.files["file"]
    if file.filename == "":
        return jsonify({"error": "No file selected"}), 400

    filename = file.filename.lower()
    text = ""

    if filename.endswith(".pdf"):
        text = extract_pdf(file.stream)
    elif filename.endswith(".docx"):
        text = extract_docx(file.stream)
    elif filename.endswith((".png", ".jpg", ".jpeg")):
        text = extract_image(file.stream)
    else:
        return jsonify({"error": "Unsupported file type"}), 400

    is_ok, details = classify_agreement(text)

    if not is_ok:
        return jsonify({
            "error": "Rejected: Not a valid agreement.",
            "details": details
        }), 400

    return jsonify({
        "filename": file.filename,
        "extracted_text": text
    })


@app.route("/export/pdf", methods=["POST"])
def export_pdf():
    text = request.form.get("text", "")
    output = io.BytesIO()
    doc = SimpleDocTemplate(output)
    styles = getSampleStyleSheet()
    story = [Paragraph(line, styles["Normal"]) for line in text.split("\n")]
    doc.build(story)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name="output.pdf")

@app.route("/export/docx", methods=["POST"])
def export_docx():
    text = request.form.get("text", "")
    output = io.BytesIO()
    d = docx.Document()
    for line in text.split("\n"):
        d.add_paragraph(line)
    d.save(output)
    output.seek(0)
    return send_file(output, as_attachment=True, download_name="output.docx")

if __name__ == "__main__":
    import os
    port = int(os.environ.get("PORT", 8000))
    app.run(host="0.0.0.0", port=port)